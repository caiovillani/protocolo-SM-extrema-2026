#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Renderizador de protocolo clinico Markdown para DOCX institucional.

Gera documento Word com qualidade editorial profissional:
- Pagina de capa institucional
- Sumario automatico (TOC)
- Tipografia Calibri com hierarquia visual
- Tabelas com header azul escuro + rows alternadas
- Fluxogramas ASCII em bloco monoespaçado
- Blockquotes com borda esquerda
- Cabecalho/rodape com paginacao
- Pagina de aprovacao/assinaturas

Uso:
    py -3.13 tools/render_protocolo_docx.py -i protocolo.md -o protocolo.docx
    py -3.13 tools/render_protocolo_docx.py -i protocolo.md -o protocolo.docx --title "Titulo"
"""

from __future__ import annotations

import argparse
import sys
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("[ERRO] python-docx nao instalado. Execute: pip install python-docx", file=sys.stderr)
    sys.exit(1)

# Import do parser compartilhado
sys.path.insert(0, str(Path(__file__).parent))
from md_protocol_parser import (
    parse_protocol, count_blocks_by_type,
    Block, Heading, Paragraph, Table, CodeBlock, ListBlock, HorizontalRule,
    Span, SpanType
)


# ===========================================================================
# CONSTANTES DE ESTILO
# ===========================================================================

# Cores institucionais — alinhadas ao template timbrado SMS Extrema
# Tema extraido de: Modelo Inicial Timbrado.docx (dk1=#0E2841, accent1=#156082)
COLOR_PRIMARY = RGBColor(14, 40, 65)        # #0E2841 azul escuro institucional (dk1)
COLOR_SECONDARY = RGBColor(21, 96, 130)     # #156082 azul petroleo (accent1)
COLOR_ACCENT = RGBColor(15, 158, 213)       # #0F9ED5 azul claro (accent4)
COLOR_GREEN = RGBColor(25, 107, 36)         # #196B24 verde institucional (accent3)
COLOR_WHITE = RGBColor(255, 255, 255)
COLOR_BLACK = RGBColor(0, 0, 0)
COLOR_GRAY_LIGHT = RGBColor(242, 242, 242)  # #F2F2F2
COLOR_GRAY_MED = RGBColor(191, 191, 191)    # #BFBFBF
COLOR_GRAY_BG = RGBColor(248, 248, 248)     # #F8F8F8

# Tipografia — alinhada ao template timbrado SMS Extrema (Aptos)
FONT_BODY = "Aptos"
FONT_HEADING = "Aptos Display"
FONT_MONO = "Consolas"

# Tamanhos
SIZE_H1 = Pt(18)
SIZE_H2 = Pt(16)
SIZE_H3 = Pt(14)
SIZE_H4 = Pt(12)
SIZE_BODY = Pt(11)
SIZE_TABLE = Pt(10)
SIZE_MONO = Pt(8.5)
SIZE_SMALL = Pt(9)

# Margens (A4) — alinhadas ao template timbrado SMS (simetrico 3.0cm)
MARGIN_TOP = Cm(2.5)
MARGIN_BOTTOM = Cm(2.5)
MARGIN_LEFT = Cm(3.0)
MARGIN_RIGHT = Cm(2.0)  # Margem menor para maximizar largura util (16cm)

# Mapa heading level -> (size, color, space_before, space_after)
HEADING_STYLES = {
    1: (SIZE_H1, COLOR_PRIMARY, Pt(24), Pt(12)),
    2: (SIZE_H2, COLOR_PRIMARY, Pt(18), Pt(10)),
    3: (SIZE_H3, COLOR_SECONDARY, Pt(14), Pt(8)),
    4: (SIZE_H4, COLOR_SECONDARY, Pt(12), Pt(6)),
}


# ===========================================================================
# HELPERS XML (python-docx low-level)
# ===========================================================================

def _set_cell_shading(cell, color: RGBColor):
    """Define cor de fundo de celula via XML."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    hex_color = f'{color[0]:02X}{color[1]:02X}{color[2]:02X}'
    shading.set(qn('w:fill'), hex_color)
    cell._element.get_or_add_tcPr().append(shading)


def _add_paragraph_border(paragraph, side: str, color: RGBColor,
                          size: int = 6, space: int = 4):
    """Adiciona borda a um paragrafo (left, right, top, bottom)."""
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = pPr.find(qn('w:pBdr'))
    if pBdr is None:
        pBdr = OxmlElement('w:pBdr')
        pPr.append(pBdr)
    border_el = OxmlElement(f'w:{side}')
    border_el.set(qn('w:val'), 'single')
    border_el.set(qn('w:sz'), str(size))
    border_el.set(qn('w:space'), str(space))
    hex_color = f'{color[0]:02X}{color[1]:02X}{color[2]:02X}'
    border_el.set(qn('w:color'), hex_color)
    pBdr.append(border_el)


def _set_cell_margins(cell, top=60, bottom=60, start=80, end=80):
    """Define margens internas da celula via XML (valores em twips/dxa).

    Defaults: top/bottom=60 (~3pt), start/end=80 (~4pt).
    """
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('bottom', bottom),
                      ('start', start), ('end', end)]:
        node = OxmlElement(f'w:{side}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


def _add_page_number_field(paragraph):
    """Insere campo PAGE no paragrafo."""
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)
    return run


# ===========================================================================
# DOCX BUILDER
# ===========================================================================

class DOCXProtocolBuilder:
    """Constroi documento DOCX institucional a partir de blocos parseados."""

    def __init__(self, metadata: dict):
        self.doc = Document()
        self.meta = metadata
        self._setup_page()

    def _setup_page(self):
        """Configura layout de pagina A4."""
        for section in self.doc.sections:
            section.page_height = Cm(29.7)
            section.page_width = Cm(21.0)
            section.top_margin = MARGIN_TOP
            section.bottom_margin = MARGIN_BOTTOM
            section.left_margin = MARGIN_LEFT
            section.right_margin = MARGIN_RIGHT

    # --- Pagina de capa ---

    def add_cover_page(self):
        """Adiciona pagina de capa institucional."""
        # Espacamento superior
        for _ in range(4):
            p = self.doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)

        # Hierarquia institucional (alinhada ao template timbrado SMS)
        p_inst = self.doc.add_paragraph()
        p_inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_inst.add_run("PREFEITURA MUNICIPAL DE EXTREMA — MG")
        run.font.name = FONT_BODY
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = COLOR_PRIMARY
        p_inst.paragraph_format.space_after = Pt(2)

        p_sms = self.doc.add_paragraph()
        p_sms.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_sms.add_run("SECRETARIA MUNICIPAL DE SAÚDE")
        run.font.name = FONT_BODY
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = COLOR_PRIMARY
        p_sms.paragraph_format.space_after = Pt(2)

        p_raps = self.doc.add_paragraph()
        p_raps.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_raps.add_run(
            "COORDENAÇÃO DA REDE DE ATENÇÃO PSICOSSOCIAL (RAPS)")
        run.font.name = FONT_BODY
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = COLOR_SECONDARY
        p_raps.paragraph_format.space_after = Pt(36)

        # Linha decorativa via paragraph border (nao Unicode)
        p_line = self.doc.add_paragraph()
        p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_line.paragraph_format.space_after = Pt(36)
        _add_paragraph_border(p_line, 'bottom', COLOR_ACCENT, size=12, space=8)

        # Titulo principal
        p_title = self.doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_text = self.meta.get(
            "title",
            "PROTOCOLO DE FLUXOS DE ATENÇÃO EM SAÚDE MENTAL")
        run = p_title.add_run(title_text)
        run.font.name = FONT_HEADING
        run.font.size = Pt(22)
        run.font.bold = True
        run.font.color.rgb = COLOR_PRIMARY
        p_title.paragraph_format.space_after = Pt(14)

        # Subtitulo
        p_sub = self.doc.add_paragraph()
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        year = self.meta.get('year', '2026')
        run = p_sub.add_run(f"Município de Extrema — MG | {year}")
        run.font.name = FONT_HEADING
        run.font.size = Pt(16)
        run.font.color.rgb = COLOR_SECONDARY
        p_sub.paragraph_format.space_after = Pt(36)

        # Linha decorativa inferior
        p_line2 = self.doc.add_paragraph()
        p_line2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_line2.paragraph_format.space_after = Pt(48)
        _add_paragraph_border(p_line2, 'bottom', COLOR_ACCENT,
                              size=12, space=8)

        # Coordenacao
        p_coord = self.doc.add_paragraph()
        p_coord.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_coord.add_run("Coordenação de Saúde Mental")
        run.font.name = FONT_BODY
        run.font.size = Pt(12)
        run.font.color.rgb = COLOR_SECONDARY
        p_coord.paragraph_format.space_after = Pt(72)

        # Data
        p_date = self.doc.add_paragraph()
        p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_text = self.meta.get("date", "Fevereiro de 2026")
        run = p_date.add_run(f"Elaborado em {date_text}")
        run.font.name = FONT_BODY
        run.font.size = SIZE_BODY
        run.font.italic = True
        run.font.color.rgb = COLOR_GRAY_MED

        self.doc.add_page_break()

    # --- Sumario automatico (TOC) ---

    def add_toc(self):
        """Insere campo TOC automatico (atualiza ao abrir no Word)."""
        p_title = self.doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p_title.add_run("SUMÁRIO")
        run.font.name = FONT_HEADING
        run.font.size = SIZE_H1
        run.font.bold = True
        run.font.color.rgb = COLOR_PRIMARY
        p_title.paragraph_format.space_after = Pt(18)

        # Campo TOC via XML
        p = self.doc.add_paragraph()
        run = p.add_run()

        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = r' TOC \o "1-4" \h \z \u '
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        # Texto placeholder visivel antes de atualizar
        placeholder_run = OxmlElement('w:r')
        placeholder_text = OxmlElement('w:t')
        placeholder_text.text = "[Clique com botão direito → Atualizar campo para gerar sumário]"
        placeholder_run.append(placeholder_text)
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run._r.append(placeholder_run)
        run._r.append(fldChar3)

        self.doc.add_page_break()

    # --- Header / Footer ---

    def setup_header_footer(self):
        """Configura cabecalho e rodape."""
        section = self.doc.sections[0]
        section.different_first_page_header_footer = True

        # --- Header (paginas normais) ---
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        title_short = self.meta.get("title_short", "Protocolo SM Extrema 2026")
        run = hp.add_run(title_short)
        run.font.name = FONT_BODY
        run.font.size = SIZE_SMALL
        run.font.italic = True
        run.font.color.rgb = COLOR_GRAY_MED

        # --- Footer (paginas normais) ---
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

        date_str = datetime.now().strftime("%d/%m/%Y")
        run = fp.add_run(f"Documento para revisão — {date_str}  |  Página ")
        run.font.name = FONT_BODY
        run.font.size = SIZE_SMALL
        run.font.italic = True
        run.font.color.rgb = COLOR_GRAY_MED

        page_run = _add_page_number_field(fp)
        page_run.font.name = FONT_BODY
        page_run.font.size = SIZE_SMALL
        page_run.font.color.rgb = COLOR_GRAY_MED

    # --- Renderizadores por tipo de bloco ---

    def render_blocks(self, blocks: list[Block]):
        """Renderiza lista de blocos sequencialmente."""
        for block in blocks:
            if isinstance(block, Heading):
                self._render_heading(block)
            elif isinstance(block, Paragraph):
                if block.style == "blockquote":
                    self._render_blockquote(block)
                else:
                    self._render_paragraph(block)
            elif isinstance(block, Table):
                self._render_table(block)
            elif isinstance(block, CodeBlock):
                self._render_code_block(block)
            elif isinstance(block, ListBlock):
                self._render_list(block)
            elif isinstance(block, HorizontalRule):
                self._render_hr()

    def _render_heading(self, h: Heading):
        """Renderiza heading com estilo Word nativo para TOC."""
        style_name = f'Heading {h.level}'
        p = self.doc.add_paragraph(style=style_name)

        # Adicionar spans
        self._add_spans_to_paragraph(p, h.spans)

        # Aplicar estilo visual customizado
        size, color, space_before, space_after = HEADING_STYLES.get(
            h.level, (SIZE_H4, COLOR_SECONDARY, Pt(12), Pt(6)))

        p.paragraph_format.space_before = space_before
        p.paragraph_format.space_after = space_after
        p.paragraph_format.keep_with_next = True
        p.paragraph_format.widow_control = True

        # H1: nova pagina (exceto primeiro heading do documento)
        if h.level == 1:
            p.paragraph_format.page_break_before = True

        for run in p.runs:
            run.font.name = FONT_HEADING
            run.font.size = size
            run.font.bold = True
            run.font.color.rgb = color

    def _render_paragraph(self, para: Paragraph):
        """Renderiza paragrafo normal."""
        p = self.doc.add_paragraph()
        self._add_spans_to_paragraph(p, para.spans)

        for run in p.runs:
            if not run.font.name:
                run.font.name = FONT_BODY
            if not run.font.size:
                run.font.size = SIZE_BODY

        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.widow_control = True

    def _render_blockquote(self, para: Paragraph):
        """Renderiza blockquote com borda esquerda azul + fundo sutil + italico."""
        p = self.doc.add_paragraph()
        self._add_spans_to_paragraph(p, para.spans)

        for run in p.runs:
            run.font.name = FONT_BODY
            run.font.size = SIZE_BODY
            run.font.italic = True

        p.paragraph_format.left_indent = Cm(1.3)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.widow_control = True

        _add_paragraph_border(p, 'left', COLOR_ACCENT, size=12, space=6)

        # Fundo azul muito sutil para destaque visual
        pPr = p._element.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'F0F5FA')
        pPr.append(shd)

    def _render_table(self, tbl: Table):
        """Renderiza tabela com estilo profissional."""
        num_cols = max(
            len(tbl.headers) if tbl.headers else 0,
            max((len(row) for row in tbl.rows), default=0)
        )
        if num_cols == 0:
            return

        num_rows = len(tbl.rows) + (1 if tbl.has_header and tbl.headers else 0)
        table = self.doc.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        table.allow_autofit = True

        row_offset = 0

        # Header row
        if tbl.has_header and tbl.headers:
            for j, header_spans in enumerate(tbl.headers):
                if j >= num_cols:
                    break
                cell = table.rows[0].cells[j]
                p = cell.paragraphs[0]
                self._add_spans_to_paragraph(p, header_spans)
                for run in p.runs:
                    run.font.name = FONT_BODY
                    run.font.size = SIZE_TABLE
                    run.font.bold = True
                    run.font.color.rgb = COLOR_WHITE
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.space_before = Pt(3)
                _set_cell_shading(cell, COLOR_PRIMARY)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                _set_cell_margins(cell)
            row_offset = 1

            # Header repete em novas paginas (via XML w:tblHeader)
            tr = table.rows[0]._tr
            trPr = tr.get_or_add_trPr()
            tbl_header = OxmlElement('w:tblHeader')
            trPr.append(tbl_header)

        # Data rows
        for i, row_data in enumerate(tbl.rows):
            doc_row_idx = i + row_offset
            if doc_row_idx >= num_rows:
                break
            for j, cell_spans in enumerate(row_data):
                if j >= num_cols:
                    break
                cell = table.rows[doc_row_idx].cells[j]
                p = cell.paragraphs[0]
                self._add_spans_to_paragraph(p, cell_spans)
                for run in p.runs:
                    if not run.font.name:
                        run.font.name = FONT_BODY
                    if not run.font.size:
                        run.font.size = SIZE_TABLE
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.space_before = Pt(2)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                _set_cell_margins(cell)

                # Rows alternadas
                if i % 2 == 0:
                    _set_cell_shading(cell, COLOR_GRAY_LIGHT)

        # Espacamento apos tabela
        spacer = self.doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(0)
        spacer.paragraph_format.space_after = Pt(8)

    def _render_code_block(self, cb: CodeBlock):
        """Renderiza bloco ASCII em tabela 1x1 monoespaçada com fundo.

        Font size se adapta a largura do conteudo para caber em A4:
        - <=65 chars: 8.5pt (normal)
        - <=80 chars: 7.5pt (reduzido)
        - >80 chars:  7.0pt (minimo)
        """
        max_line_len = max((len(line) for line in cb.lines), default=0)

        # Font size adaptativo
        if max_line_len > 80:
            font_size = Pt(7.0)
        elif max_line_len > 65:
            font_size = Pt(7.5)
        else:
            font_size = SIZE_MONO  # Pt(8.5)

        table = self.doc.add_table(rows=1, cols=1)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False

        cell = table.rows[0].cells[0]
        cell.width = Cm(16.0)  # Largura util da pagina (21 - 3 - 2 = 16cm)
        _set_cell_shading(cell, COLOR_GRAY_BG)
        _set_cell_margins(cell, top=60, bottom=60, start=80, end=80)

        p = cell.paragraphs[0]
        content = '\n'.join(cb.lines)
        run = p.add_run(content)
        run.font.name = FONT_MONO
        run.font.size = font_size
        run.font.color.rgb = COLOR_BLACK

        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)

        # Aviso se conteudo pode estourar largura
        if max_line_len > 90:
            print(f"    [WARN] Code block com {max_line_len} chars — "
                  f"pode exceder largura util a {font_size.pt}pt")

        # Espacamento apos
        spacer = self.doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(0)
        spacer.paragraph_format.space_after = Pt(8)

    def _render_list(self, lb: ListBlock):
        """Renderiza lista ordenada ou nao-ordenada."""
        for idx, item_spans in enumerate(lb.items, start=lb.start):
            style = 'List Number' if lb.ordered else 'List Bullet'
            p = self.doc.add_paragraph(style=style)
            self._add_spans_to_paragraph(p, item_spans)
            for run in p.runs:
                if not run.font.name:
                    run.font.name = FONT_BODY
                if not run.font.size:
                    run.font.size = SIZE_BODY
            p.paragraph_format.space_after = Pt(3)

    def _render_hr(self):
        """Renderiza linha horizontal como borda inferior."""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        _add_paragraph_border(p, 'bottom', COLOR_GRAY_MED, size=12, space=4)

    # --- Pagina de aprovacao ---

    def add_approval_page(self):
        """Adiciona pagina final com tabela de assinaturas."""
        self.doc.add_page_break()

        p_title = self.doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_title.add_run("APROVAÇÃO")
        run.font.name = FONT_HEADING
        run.font.size = SIZE_H1
        run.font.bold = True
        run.font.color.rgb = COLOR_PRIMARY
        p_title.paragraph_format.space_after = Pt(24)

        p_intro = self.doc.add_paragraph()
        run = p_intro.add_run(
            "Este protocolo foi elaborado pela Coordenação de Saúde Mental "
            "da Secretaria Municipal de Saúde de Extrema e aprovado "
            "em ___/___/______."
        )
        run.font.name = FONT_BODY
        run.font.size = SIZE_BODY
        p_intro.paragraph_format.space_after = Pt(24)

        # Tabela de assinaturas
        cargos = [
            "Secretário(a) Municipal de Saúde",
            "Coordenador(a) de Saúde Mental",
            "Coordenador(a) da Atenção Primária",
        ]

        table = self.doc.add_table(rows=len(cargos) + 1, cols=3)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header
        headers = ["Cargo", "Nome", "Assinatura e Data"]
        for j, h_text in enumerate(headers):
            cell = table.rows[0].cells[j]
            p = cell.paragraphs[0]
            run = p.add_run(h_text)
            run.font.name = FONT_BODY
            run.font.size = SIZE_TABLE
            run.font.bold = True
            run.font.color.rgb = COLOR_WHITE
            _set_cell_shading(cell, COLOR_PRIMARY)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_cell_margins(cell)

        # Linhas de assinatura com guias pontilhados
        for i, cargo in enumerate(cargos, start=1):
            cell_cargo = table.rows[i].cells[0]
            p = cell_cargo.paragraphs[0]
            run = p.add_run(cargo)
            run.font.name = FONT_BODY
            run.font.size = SIZE_TABLE
            cell_cargo.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_cell_margins(cell_cargo)

            # Celulas Nome e Assinatura com borda inferior pontilhada
            for j in range(1, 3):
                cell = table.rows[i].cells[j]
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                _set_cell_margins(cell)
                cell.paragraphs[0].paragraph_format.space_after = Pt(24)
                # Borda inferior pontilhada como guia de assinatura
                tc = cell._element
                tcPr = tc.get_or_add_tcPr()
                tcBorders = OxmlElement('w:tcBorders')
                bottom_border = OxmlElement('w:bottom')
                bottom_border.set(qn('w:val'), 'dotted')
                bottom_border.set(qn('w:sz'), '4')
                bottom_border.set(qn('w:space'), '0')
                bottom_border.set(qn('w:color'), 'BFBFBF')
                tcBorders.append(bottom_border)
                tcPr.append(tcBorders)

        # Rodape do documento
        spacer = self.doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(48)

        p_footer = self.doc.add_paragraph()
        p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_text = self.meta.get("date", "fevereiro de 2026")
        run = p_footer.add_run(f"Documento elaborado em {date_text}")
        run.font.name = FONT_BODY
        run.font.size = SIZE_SMALL
        run.font.italic = True
        run.font.color.rgb = COLOR_GRAY_MED

        p_org = self.doc.add_paragraph()
        p_org.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_org.add_run("Secretaria Municipal de Saúde — Extrema/MG")
        run.font.name = FONT_BODY
        run.font.size = SIZE_SMALL
        run.font.italic = True
        run.font.color.rgb = COLOR_GRAY_MED

    # --- Helpers inline ---

    def _add_spans_to_paragraph(self, paragraph, spans: list[Span]):
        """Adiciona lista de Spans como runs formatados ao paragrafo."""
        for span in spans:
            run = paragraph.add_run(span.text)

            if span.type == SpanType.BOLD:
                run.font.bold = True
            elif span.type == SpanType.ITALIC:
                run.font.italic = True
            elif span.type == SpanType.EMOJI:
                run.font.name = 'Segoe UI Emoji'
            elif span.type == SpanType.LINK:
                run.font.color.rgb = COLOR_SECONDARY
                run.font.underline = True

    # --- Salvar ---

    def save(self, output_path: str):
        """Salva documento DOCX."""
        self.doc.save(output_path)


# ===========================================================================
# CLI
# ===========================================================================

def main():
    parser = argparse.ArgumentParser(
        description="Renderiza protocolo clinico Markdown para DOCX institucional."
    )
    parser.add_argument("-i", "--input", required=True, type=Path,
                        help="Arquivo Markdown de entrada")
    parser.add_argument("-o", "--output", required=True, type=Path,
                        help="Arquivo DOCX de saida")
    parser.add_argument("--title",
                        default="PROTOCOLO DE FLUXOS DE ATENÇÃO EM SAÚDE MENTAL",
                        help="Titulo do documento (capa)")
    parser.add_argument("--title-short",
                        default="Protocolo SM Extrema 2026",
                        help="Titulo abreviado (cabecalho)")
    parser.add_argument("--year", default="2026",
                        help="Ano de elaboracao")
    parser.add_argument("--date", default="fevereiro de 2026",
                        help="Data de elaboracao por extenso")

    args = parser.parse_args()

    # Validacao
    if not args.input.exists():
        print(f"[ERRO] Arquivo nao encontrado: {args.input}", file=sys.stderr)
        sys.exit(1)

    # Ler e parsear
    print(f"[.] Lendo {args.input.name}...")
    md_content = args.input.read_text(encoding='utf-8')

    print(f"[.] Parseando Markdown...")
    blocks = parse_protocol(md_content)
    counts = count_blocks_by_type(blocks)

    print(f"    {len(blocks)} blocos: {', '.join(f'{k}={v}' for k, v in sorted(counts.items()))}")

    # Construir DOCX
    print(f"[.] Construindo DOCX...")
    metadata = {
        "title": args.title,
        "title_short": args.title_short,
        "year": args.year,
        "date": args.date,
    }

    builder = DOCXProtocolBuilder(metadata)
    builder.add_cover_page()
    builder.add_toc()
    builder.setup_header_footer()
    builder.render_blocks(blocks)
    builder.add_approval_page()

    # Salvar
    builder.save(str(args.output))

    size_kb = args.output.stat().st_size / 1024
    print(f"[OK] Documento gerado: {args.output}")
    print(f"     Tamanho: {size_kb:.1f} KB")
    print()
    print("[!] LEMBRETE: Ao abrir no Word, atualize o sumario:")
    print("    Clique com botao direito no sumario -> 'Atualizar campo' -> 'Atualizar pagina inteira'")


if __name__ == "__main__":
    main()
