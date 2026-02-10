#!/usr/bin/env python3
"""
CRIE Document Extraction and Organization Tool

Extrai, transcreve e organiza 27 documentos do acervo Licitacoes/CRIE
em estrutura tematica padronizada com frontmatter YAML.

Formatos suportados: PDF, DOCX, MD, TXT
Saida: Markdown com frontmatter YAML + source_manifest.yaml + indice

Usage:
    py -3.13 tools/extract_crie_batch.py \\
        --source "path/to/Licitacoes_Compras_e_Contratos/" \\
        --output "referencias_crie/" \\
        [--dry-run]
"""

import argparse
import hashlib
import io
import json
import sys
import yaml
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional

# Fix Windows console encoding for UTF-8 filenames
if sys.platform == "win32":
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')

try:
    import pdfplumber
except ImportError:
    print("ERRO: pdfplumber nao instalado. Execute: pip install pdfplumber", file=sys.stderr)
    sys.exit(1)

try:
    from docx import Document as DocxDocument
except ImportError:
    print("ERRO: python-docx nao instalado. Execute: pip install python-docx", file=sys.stderr)
    sys.exit(1)


# ===== FILE CONFIGURATION =====

FILE_CONFIGS: List[Dict[str, Any]] = [
    # --- 01_normativos_mrosc (7 files) ---
    {
        "src": "normativos/DECRETO Nº 47.132, DE 20 DE JANEIRO.txt",
        "dst": "01_normativos_mrosc/DECRETO_MG_47132_2017.md",
        "doc_id": "CRIE-NOR-001",
        "title": "Decreto MG n. 47.132/2017 - Regulamenta parcerias com OSCs (MROSC estadual)",
        "doc_type": "decreto",
        "taxonomy": "normativos_mrosc",
        "year": 2017,
        "publisher": "Governo do Estado de Minas Gerais",
        "topics": ["MROSC", "parcerias", "OSC", "termo de colaboracao", "termo de fomento"],
    },
    {
        "src": "normativos/DELIBERAÇÃO_CIB_SUS_MG_1403_2013.pdf",
        "dst": "01_normativos_mrosc/DELIBERACAO_CIB_SUS_MG_1403_2013.md",
        "doc_id": "CRIE-NOR-002",
        "title": "Deliberacao CIB-SUS/MG n. 1.403/2013",
        "doc_type": "resolucao",
        "taxonomy": "normativos_mrosc",
        "year": 2013,
        "publisher": "CIB-SUS/MG",
        "topics": ["SUS", "governanca", "saude", "CIB", "Minas Gerais"],
    },
    {
        "src": "normativos/L12764.pdf",
        "dst": "01_normativos_mrosc/LEI_12764_2012_BERENICE_PIANA.md",
        "doc_id": "CRIE-NOR-003",
        "title": "Lei Federal n. 12.764/2012 - Politica Nacional de Protecao dos Direitos da Pessoa com TEA",
        "doc_type": "legislacao",
        "taxonomy": "normativos_mrosc",
        "year": 2012,
        "publisher": "Presidencia da Republica",
        "topics": ["TEA", "autismo", "direitos", "protecao", "Berenice Piana"],
    },
    {
        "src": "normativos/Lei Entidades Beneficentes.pdf",
        "dst": "01_normativos_mrosc/LEI_ENTIDADES_BENEFICENTES.md",
        "doc_id": "CRIE-NOR-004",
        "title": "Lei Federal - Entidades Beneficentes de Assistencia Social",
        "doc_type": "legislacao",
        "taxonomy": "normativos_mrosc",
        "year": None,
        "publisher": "Presidencia da Republica",
        "topics": ["entidades beneficentes", "assistencia social", "CEBAS", "imunidade tributaria"],
    },
    {
        "src": "normativos/licitacoes_e_contratos_administrativos_L14133.pdf",
        "dst": "01_normativos_mrosc/LEI_14133_2021_LICITACOES.md",
        "doc_id": "CRIE-NOR-005",
        "title": "Lei Federal n. 14.133/2021 - Nova Lei de Licitacoes e Contratos Administrativos",
        "doc_type": "legislacao",
        "taxonomy": "normativos_mrosc",
        "year": 2021,
        "publisher": "Presidencia da Republica",
        "topics": ["licitacoes", "contratos", "compras publicas", "pregao", "concorrencia"],
    },
    {
        "src": "normativos/OSC_Leis, Decretos e Normas do Ministério da Saúde.pdf",
        "dst": "01_normativos_mrosc/COMPILACAO_OSC_NORMAS_MS.md",
        "doc_id": "CRIE-NOR-006",
        "title": "Compilacao - Leis, Decretos e Normas do MS sobre OSCs",
        "doc_type": "compilacao",
        "taxonomy": "normativos_mrosc",
        "year": None,
        "publisher": "Ministerio da Saude",
        "topics": ["OSC", "normas MS", "parcerias saude", "MROSC saude"],
    },
    {
        "src": "normativos/PORTARIA INTERMINISTERIAL SG_MGI_AGU Nº 197, DE 11 DE AGOSTO DE 2025 — Transferegov.br.pdf",
        "dst": "01_normativos_mrosc/PORTARIA_197_2025_TRANSFEREGOV.md",
        "doc_id": "CRIE-NOR-007",
        "title": "Portaria Interministerial SG/MGI/AGU n. 197/2025 - Transferegov.br",
        "doc_type": "portaria",
        "taxonomy": "normativos_mrosc",
        "year": 2025,
        "publisher": "SG/MGI/AGU",
        "topics": ["transferencias", "Transferegov", "convenios", "parcerias federais"],
    },

    # --- 02_projeto_crie (4 files) ---
    {
        "src": "projeto_crie/CUSTO FUNCIONARIO com educador fisico 2.pdf",
        "dst": "02_projeto_crie/CUSTO_FUNCIONARIO_CRIE.md",
        "doc_id": "CRIE-PRJ-001",
        "title": "Planilha de Custos - Funcionarios CRIE (com Educador Fisico)",
        "doc_type": "planilha",
        "taxonomy": "projeto_crie",
        "year": 2025,
        "publisher": "CRIE/APAE Extrema",
        "topics": ["custos", "RH", "encargos", "fisioterapeuta", "psicologo", "educador fisico"],
    },
    {
        "src": "projeto_crie/PLANO DE TRABALHO CRIE.docx",
        "dst": "02_projeto_crie/PLANO_TRABALHO_CRIE.md",
        "doc_id": "CRIE-PRJ-002",
        "title": "Plano de Trabalho - Parceria SMS Extrema x CRIE",
        "doc_type": "plano_trabalho",
        "taxonomy": "projeto_crie",
        "year": 2026,
        "publisher": "CRIE - Centro de Integracao Especial",
        "topics": ["plano de trabalho", "parceria", "MROSC", "CRIE", "PCD", "TEA"],
    },
    {
        "src": "projeto_crie/RELATÓRIO_TÉCNICO_CAPS_I e II.docx",
        "dst": "02_projeto_crie/RELATORIO_TECNICO_CAPS_I_II.md",
        "doc_id": "CRIE-PRJ-003",
        "title": "Relatorio Tecnico - CAPS I e II Extrema",
        "doc_type": "relatorio_tecnico",
        "taxonomy": "projeto_crie",
        "year": 2025,
        "publisher": "SMS Extrema",
        "topics": ["CAPS", "saude mental", "capacidade", "demanda", "parceria CRIE"],
    },
    {
        "src": "projeto_crie/TR - CRIE.docx",
        "dst": "02_projeto_crie/TR_CRIE.md",
        "doc_id": "CRIE-PRJ-004",
        "title": "Termo de Referencia - Parceria SMS Extrema x CRIE",
        "doc_type": "termo_referencia",
        "taxonomy": "projeto_crie",
        "year": 2026,
        "publisher": "SMS Extrema",
        "topics": ["termo de referencia", "parceria", "MROSC", "dispensa", "PCD", "TEA"],
    },

    # --- 03_sinteses_analise (11 files) ---
    {
        "src": "sinteses_ia/BANCO_CONHECIMENTO_TECNICO.md",
        "dst": "03_sinteses_analise/BANCO_CONHECIMENTO_TECNICO.md",
        "doc_id": "CRIE-SIN-001",
        "title": "Banco de Conhecimento Tecnico - Parceria CRIE",
        "doc_type": "analise",
        "taxonomy": "sinteses_analise",
        "year": 2026,
        "publisher": "Analise Tecnica IA",
        "topics": ["base legal", "custos", "MROSC", "vedacoes", "capacidade CRIE"],
    },
    {
        "src": "sinteses_ia/CONTEXTO_ESTRATEGICO_EXTREMA.md",
        "dst": "03_sinteses_analise/CONTEXTO_ESTRATEGICO_EXTREMA.md",
        "doc_id": "CRIE-SIN-002",
        "title": "Contexto Estrategico - Parceria CRIE x SMS Extrema",
        "doc_type": "analise",
        "taxonomy": "sinteses_analise",
        "year": 2026,
        "publisher": "Analise Tecnica IA",
        "topics": ["contexto", "riscos", "validacao", "parceria CRIE"],
    },
    {
        "src": "sinteses_ia/EXTRACAO_DADOS_CRIE.md",
        "dst": "03_sinteses_analise/EXTRACAO_DADOS_CRIE.md",
        "doc_id": "CRIE-SIN-003",
        "title": "Extracao de Dados - Documentos CRIE",
        "doc_type": "extracao",
        "taxonomy": "sinteses_analise",
        "year": 2026,
        "publisher": "Analise Tecnica IA",
        "topics": ["extracao", "dados", "TR", "plano trabalho", "CRIE"],
    },
    {
        "src": "sinteses_ia/LISTA_EXIGENCIAS_COMPLEMENTACAO.md",
        "dst": "03_sinteses_analise/LISTA_EXIGENCIAS_COMPLEMENTACAO.md",
        "doc_id": "CRIE-SIN-004",
        "title": "Lista de Exigencias de Complementacao - CRIE",
        "doc_type": "checklist",
        "taxonomy": "sinteses_analise",
        "year": 2026,
        "publisher": "Analise Tecnica IA",
        "topics": ["exigencias", "documentacao", "CAGEC", "certidoes", "compliance"],
    },
    {
        "src": "sinteses_ia/LOG_AUDITORIA_TAREFAS.md",
        "dst": "03_sinteses_analise/LOG_AUDITORIA_TAREFAS.md",
        "doc_id": "CRIE-SIN-005",
        "title": "Log de Auditoria de Tarefas - Projeto CRIE",
        "doc_type": "log",
        "taxonomy": "sinteses_analise",
        "year": 2026,
        "publisher": "Analise Tecnica IA",
        "topics": ["auditoria", "tarefas", "progresso", "fases"],
    },
    {
        "src": "sinteses_ia/MEMORIA_CALCULO_ORCAMENTO.md",
        "dst": "03_sinteses_analise/MEMORIA_CALCULO_ORCAMENTO.md",
        "doc_id": "CRIE-SIN-006",
        "title": "Memoria de Calculo Orcamentario - Parceria CRIE",
        "doc_type": "analise",
        "taxonomy": "sinteses_analise",
        "year": 2026,
        "publisher": "Analise Tecnica IA",
        "topics": ["orcamento", "custos", "INSS", "FGTS", "encargos"],
    },
    {
        "src": "sinteses_ia/planejamento-executivo-diretoria_contrato_compras_licitacoes-extrema.md",
        "dst": "03_sinteses_analise/PLANEJAMENTO_EXECUTIVO_DIRETORIA.md",
        "doc_id": "CRIE-SIN-007",
        "title": "Planejamento Executivo - Diretoria de Contratos e Licitacoes SMS Extrema",
        "doc_type": "planejamento",
        "taxonomy": "sinteses_analise",
        "year": 2026,
        "publisher": "Analise Tecnica IA",
        "topics": ["planejamento", "licitacoes", "contratos", "L14133", "governanca", "KPIs"],
    },
    {
        "src": "sinteses_ia/RASCUNHO_INSTRUMENTOS_JURIDICOS.md",
        "dst": "03_sinteses_analise/RASCUNHO_INSTRUMENTOS_JURIDICOS.md",
        "doc_id": "CRIE-SIN-008",
        "title": "Rascunho de Instrumentos Juridicos - TR e PT Parceria CRIE",
        "doc_type": "minuta",
        "taxonomy": "sinteses_analise",
        "year": 2026,
        "publisher": "Analise Tecnica IA",
        "topics": ["termo referencia", "plano trabalho", "minuta", "MROSC", "Decreto MG 47.132"],
    },
    {
        "src": "sinteses_ia/RELATORIO_ENTREGA_DOCUMENTOS_OFICIAIS.md",
        "dst": "03_sinteses_analise/RELATORIO_ENTREGA_DOCUMENTOS.md",
        "doc_id": "CRIE-SIN-009",
        "title": "Relatorio de Entrega de Documentos Oficiais - Projeto CRIE",
        "doc_type": "relatorio",
        "taxonomy": "sinteses_analise",
        "year": 2026,
        "publisher": "Analise Tecnica IA",
        "topics": ["entrega", "documentos", "oficiais", "CRIE"],
    },
    {
        "src": "sinteses_ia/RELATORIO_FINAL_ENTREGA.md",
        "dst": "03_sinteses_analise/RELATORIO_FINAL_ENTREGA.md",
        "doc_id": "CRIE-SIN-010",
        "title": "Relatorio Final de Entrega - Projeto CRIE",
        "doc_type": "relatorio",
        "taxonomy": "sinteses_analise",
        "year": 2026,
        "publisher": "Analise Tecnica IA",
        "topics": ["entrega final", "qualidade", "conformidade", "proximos passos"],
    },
    {
        "src": "sinteses_ia/RELATORIO_NAO_CONFORMIDADES.md",
        "dst": "03_sinteses_analise/RELATORIO_NAO_CONFORMIDADES.md",
        "doc_id": "CRIE-SIN-011",
        "title": "Relatorio de Nao-Conformidades - Parceria CRIE",
        "doc_type": "analise_risco",
        "taxonomy": "sinteses_analise",
        "year": 2026,
        "publisher": "Analise Tecnica IA",
        "topics": ["riscos", "nao-conformidades", "dispensa", "CAGEC", "objeto hibrido"],
    },

    # --- 04_benchmarks (5 files) ---
    {
        "src": "benchmarks/exemplo_apae_pouso_alegre.pdf",
        "dst": "04_benchmarks/APAE_POUSO_ALEGRE_INFRAESTRUTURA.md",
        "doc_id": "CRIE-BEN-001",
        "title": "Benchmark - Plano de Trabalho Instituto Filippo Smaldone (Pouso Alegre) - Infraestrutura",
        "doc_type": "benchmark",
        "taxonomy": "benchmarks",
        "year": 2023,
        "publisher": "Instituto Filippo Smaldone / Pouso Alegre-MG",
        "topics": ["APAE", "infraestrutura", "acessibilidade", "SCFV", "PCD"],
    },
    {
        "src": "benchmarks/exemplo_apae_2.pdf",
        "dst": "04_benchmarks/AVANTE_SOCIAL_SCFV_POUSO_ALEGRE.md",
        "doc_id": "CRIE-BEN-002",
        "title": "Benchmark - Plano de Trabalho Instituto Avante Social - SCFV Pouso Alegre",
        "doc_type": "benchmark",
        "taxonomy": "benchmarks",
        "year": 2025,
        "publisher": "Instituto Avante Social / Pouso Alegre-MG",
        "topics": ["SCFV", "assistencia social", "720 usuarios", "4 centros", "24 meses"],
    },
    {
        "src": "benchmarks/copy_of_ANEXO_IV_MODELO_DE_PLANO_DE_TRABALHO_ajustada_27.09.2023.pdf",
        "dst": "04_benchmarks/MODELO_PLANO_TRABALHO_MINC.md",
        "doc_id": "CRIE-BEN-003",
        "title": "Template - Anexo IV Modelo de Plano de Trabalho (Ministerio da Cultura)",
        "doc_type": "template",
        "taxonomy": "benchmarks",
        "year": 2023,
        "publisher": "Ministerio da Cultura",
        "topics": ["template", "plano trabalho", "OSC", "cultura", "modelo"],
    },
    {
        "src": "benchmarks/Estudo+Tecnico+Preliminar+000004+2025.pdf",
        "dst": "04_benchmarks/ETP_000004_2025.md",
        "doc_id": "CRIE-BEN-004",
        "title": "Estudo Tecnico Preliminar n. 000004/2025",
        "doc_type": "etp",
        "taxonomy": "benchmarks",
        "year": 2025,
        "publisher": "Administracao Municipal",
        "topics": ["ETP", "estudo tecnico", "licitacao", "planejamento"],
    },
    {
        "src": "benchmarks/Plano_de_trabalho_Saude_para_manutencao_de_pagamento_e_educador_fisico_assinado.pdf",
        "dst": "04_benchmarks/PLANO_TRABALHO_SAUDE_EDUC_FISICO.md",
        "doc_id": "CRIE-BEN-005",
        "title": "Plano de Trabalho - Saude (Manutencao e Educador Fisico)",
        "doc_type": "plano_trabalho",
        "taxonomy": "benchmarks",
        "year": 2025,
        "publisher": "Administracao Municipal",
        "topics": ["plano trabalho", "saude", "educador fisico", "manutencao"],
    },
]


# ===== UTILITIES =====

def compute_md5(file_path: Path) -> str:
    """Calcula MD5 hash do arquivo."""
    md5_hash = hashlib.md5()
    try:
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(8192), b""):
                md5_hash.update(chunk)
        return md5_hash.hexdigest()
    except Exception as e:
        print(f"  [!] MD5 falhou para {file_path.name}: {e}", file=sys.stderr)
        return "error"


def generate_frontmatter(config: Dict, file_hash: str, pages: Optional[int],
                         original_format: str, extraction_method: str) -> str:
    """Gera bloco YAML frontmatter."""
    fm = {
        "doc_id": config["doc_id"],
        "title": config["title"],
        "taxonomy": config["taxonomy"],
        "doc_type": config["doc_type"],
        "source_file": config["src"],
        "original_format": original_format,
        "language": "pt-BR",
        "publisher": config["publisher"],
        "topics": config["topics"],
        "related_documents": [],
        "file_hash": file_hash,
        "extraction_method": extraction_method,
        "extraction_date": datetime.now().strftime("%Y-%m-%d"),
        "needs_human_review": False,
    }
    if config.get("year"):
        fm["year"] = config["year"]
    if pages:
        fm["pages_total"] = pages
    return "---\n" + yaml.dump(fm, allow_unicode=True, default_flow_style=False, sort_keys=False) + "---\n\n"


# ===== EXTRACTORS =====

def extract_pdf(file_path: Path) -> tuple[str, int]:
    """Extrai texto de PDF via pdfplumber. Retorna (texto, num_paginas)."""
    pages_text = []
    try:
        with pdfplumber.open(file_path) as pdf:
            num_pages = len(pdf.pages)
            for i, page in enumerate(pdf.pages):
                text = page.extract_text() or ""
                if text.strip():
                    pages_text.append(text)
                # Tables
                tables = page.extract_tables()
                for table in tables:
                    if table:
                        md_table = _pdf_table_to_markdown(table)
                        if md_table:
                            pages_text.append(md_table)
    except Exception as e:
        print(f"  [X] Erro PDF {file_path.name}: {e}", file=sys.stderr)
        return "", 0
    return "\n\n".join(pages_text), num_pages


def _pdf_table_to_markdown(table: list) -> str:
    """Converte tabela pdfplumber (lista de listas) para markdown pipe-table."""
    if not table or len(table) < 2:
        return ""
    # Clean cells
    clean = []
    for row in table:
        clean_row = [(cell or "").replace("\n", " ").strip() for cell in row]
        clean.append(clean_row)

    if not clean:
        return ""

    # Header
    header = clean[0]
    num_cols = len(header)
    lines = ["| " + " | ".join(header) + " |"]
    lines.append("| " + " | ".join(["---"] * num_cols) + " |")
    for row in clean[1:]:
        # Pad row to match header length
        padded = row + [""] * (num_cols - len(row))
        lines.append("| " + " | ".join(padded[:num_cols]) + " |")
    return "\n".join(lines)


def extract_docx(file_path: Path) -> str:
    """Extrai texto de DOCX preservando ordem de paragrafos e tabelas."""
    try:
        doc = DocxDocument(str(file_path))
    except Exception as e:
        print(f"  [X] Erro DOCX {file_path.name}: {e}", file=sys.stderr)
        return ""

    # Build lookup maps: element id -> python-docx wrapper
    para_map = {id(p._element): p for p in doc.paragraphs}
    table_map = {id(t._element): t for t in doc.tables}

    lines = []
    for child in doc.element.body:
        child_id = id(child)

        if child_id in para_map:
            para = para_map[child_id]
            text = para.text.strip()
            if not text:
                continue

            style_name = para.style.name or ""

            if "Heading" in style_name:
                level = _extract_heading_level(style_name)
                lines.append(f"{'#' * level} {text}")
            elif style_name.startswith("List"):
                if "Number" in style_name or "Ordered" in style_name:
                    lines.append(f"1. {text}")
                else:
                    lines.append(f"- {text}")
            else:
                formatted = _format_docx_runs(para)
                lines.append(formatted)

        elif child_id in table_map:
            table = table_map[child_id]
            md_table = _docx_table_to_markdown(table)
            if md_table:
                lines.append(md_table)

    return "\n\n".join(lines)


def _extract_heading_level(style_name: str) -> int:
    """Extrai nivel de heading do nome do estilo Word."""
    for word in style_name.split():
        if word.isdigit():
            return min(int(word), 6)
    return 1


def _format_docx_runs(para) -> str:
    """Formata runs de paragrafo DOCX com bold/italic markdown."""
    if not para.runs:
        return para.text

    parts = []
    for run in para.runs:
        text = run.text
        if not text:
            continue
        if run.bold and run.italic:
            parts.append(f"***{text}***")
        elif run.bold:
            parts.append(f"**{text}**")
        elif run.italic:
            parts.append(f"*{text}*")
        else:
            parts.append(text)

    return "".join(parts) if parts else para.text


def _docx_table_to_markdown(table) -> str:
    """Converte tabela python-docx para markdown pipe-table."""
    rows_data = []
    for row in table.rows:
        cells = [cell.text.replace("\n", " ").strip() for cell in row.cells]
        rows_data.append(cells)

    if len(rows_data) < 1:
        return ""

    num_cols = max(len(r) for r in rows_data)
    header = rows_data[0]
    header += [""] * (num_cols - len(header))

    lines = ["| " + " | ".join(header) + " |"]
    lines.append("| " + " | ".join(["---"] * num_cols) + " |")

    for row in rows_data[1:]:
        padded = row + [""] * (num_cols - len(row))
        lines.append("| " + " | ".join(padded[:num_cols]) + " |")

    return "\n".join(lines)


def extract_md(file_path: Path) -> str:
    """Le arquivo Markdown (copia conteudo)."""
    try:
        return file_path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        return file_path.read_text(encoding="latin-1")


def extract_txt(file_path: Path) -> str:
    """Le arquivo texto."""
    try:
        return file_path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        return file_path.read_text(encoding="latin-1")


# ===== MAIN PROCESSOR =====

def process_file(config: Dict, source_dir: Path, output_dir: Path,
                 dry_run: bool = False) -> Dict[str, Any]:
    """Processa um arquivo: extrai, gera frontmatter, salva."""
    src_path = source_dir / config["src"]
    dst_path = output_dir / config["dst"]

    result = {
        "doc_id": config["doc_id"],
        "src": config["src"],
        "dst": config["dst"],
        "status": "pending",
        "pages": None,
        "file_hash": None,
        "chars_extracted": 0,
    }

    # Check source exists
    if not src_path.exists():
        result["status"] = "source_not_found"
        print(f"  [X] NAO ENCONTRADO: {config['src']}", file=sys.stderr)
        return result

    # Compute hash
    file_hash = compute_md5(src_path)
    result["file_hash"] = file_hash

    if dry_run:
        ext = src_path.suffix.lower()
        result["status"] = f"dry_run ({ext})"
        print(f"  [~] {config['doc_id']} | {config['src']} -> {config['dst']}")
        return result

    # Extract based on format
    ext = src_path.suffix.lower()
    pages = None

    if ext == ".pdf":
        content, pages = extract_pdf(src_path)
        original_format = "PDF"
        extraction_method = "pdfplumber"
        result["pages"] = pages
    elif ext == ".docx":
        content = extract_docx(src_path)
        original_format = "DOCX"
        extraction_method = "python-docx"
    elif ext == ".md":
        content = extract_md(src_path)
        original_format = "MD"
        extraction_method = "copy"
    elif ext == ".txt":
        content = extract_txt(src_path)
        original_format = "TXT"
        extraction_method = "read"
    else:
        result["status"] = f"unsupported_format ({ext})"
        print(f"  [X] Formato nao suportado: {ext} para {config['src']}")
        return result

    if not content.strip():
        result["status"] = "empty_extraction"
        print(f"  [!] Extracao vazia: {config['src']}")
        return result

    result["chars_extracted"] = len(content)

    # Generate frontmatter
    frontmatter = generate_frontmatter(
        config=config,
        file_hash=file_hash,
        pages=pages,
        original_format=original_format,
        extraction_method=extraction_method,
    )

    # For MD files that already have frontmatter, strip existing before prepending new
    if ext == ".md" and content.startswith("---"):
        # Find end of existing frontmatter
        end_idx = content.find("---", 3)
        if end_idx != -1:
            content = content[end_idx + 3:].lstrip("\n")

    # Build final document
    doc_title = f"# {config['title']}\n\n"
    # Only add title heading if content doesn't already start with one
    if content.lstrip().startswith("# "):
        final_content = frontmatter + content
    else:
        final_content = frontmatter + doc_title + content

    # Write output
    dst_path.parent.mkdir(parents=True, exist_ok=True)
    dst_path.write_text(final_content, encoding="utf-8")

    result["status"] = "success"
    print(f"  [>] {config['doc_id']} | {dst_path.name} ({len(content):,} chars)")
    return result


def generate_manifest(results: List[Dict], output_dir: Path):
    """Gera source_manifest.yaml."""
    manifest = {
        "metadata": {
            "generated_at": datetime.now().isoformat(),
            "total_documents": len(results),
            "generator": "tools/extract_crie_batch.py",
            "project": "Protocolo SM Extrema 2026 - Acervo CRIE/Licitacoes",
        },
        "summary": {
            "by_status": {},
            "by_format": {},
            "total_chars": sum(r.get("chars_extracted", 0) for r in results),
        },
        "documents": results,
    }

    # Compute summaries
    for r in results:
        status = r.get("status", "unknown")
        manifest["summary"]["by_status"][status] = manifest["summary"]["by_status"].get(status, 0) + 1
        ext = Path(r["src"]).suffix.lower()
        manifest["summary"]["by_format"][ext] = manifest["summary"]["by_format"].get(ext, 0) + 1

    manifest_path = output_dir / "_config" / "source_manifest.yaml"
    manifest_path.parent.mkdir(parents=True, exist_ok=True)
    manifest_path.write_text(
        yaml.dump(manifest, allow_unicode=True, default_flow_style=False, sort_keys=False),
        encoding="utf-8",
    )
    print(f"\n[>] Manifesto: {manifest_path}")


def generate_index_md(results: List[Dict], output_dir: Path):
    """Gera 00_INDICE_CRIE.md com tabela legivel."""
    lines = [
        "# Indice — Acervo CRIE / Licitacoes e Contratos",
        "",
        f"**Gerado:** {datetime.now().strftime('%Y-%m-%d')}",
        f"**Total documentos:** {len(results)}",
        "",
        "---",
        "",
    ]

    # Group by taxonomy folder
    groups = {}
    for r in results:
        folder = Path(r["dst"]).parts[0]
        if folder not in groups:
            groups[folder] = []
        groups[folder].append(r)

    folder_labels = {
        "01_normativos_mrosc": "Normativos MROSC / Legislacao",
        "02_projeto_crie": "Projeto CRIE — Documentos Originais",
        "03_sinteses_analise": "Sinteses e Analises Tecnicas",
        "04_benchmarks": "Benchmarks / Referencias Externas",
    }

    for folder in sorted(groups.keys()):
        label = folder_labels.get(folder, folder)
        items = groups[folder]
        lines.append(f"## {label}")
        lines.append("")
        lines.append("| doc_id | Titulo | Formato | Status |")
        lines.append("|--------|--------|---------|--------|")
        for r in items:
            ext = Path(r["src"]).suffix.upper().lstrip(".")
            # Find config for title
            cfg = next((c for c in FILE_CONFIGS if c["doc_id"] == r["doc_id"]), None)
            title = cfg["title"] if cfg else r["doc_id"]
            lines.append(f"| {r['doc_id']} | {title} | {ext} | {r['status']} |")
        lines.append("")

    index_path = output_dir / "00_INDICE_CRIE.md"
    index_path.write_text("\n".join(lines), encoding="utf-8")
    print(f"[>] Indice: {index_path}")


# ===== CLI =====

def main():
    parser = argparse.ArgumentParser(
        description="Extrai e organiza documentos CRIE/Licitacoes em Markdown padronizado."
    )
    parser.add_argument(
        "--source", type=Path, required=True,
        help="Diretorio fonte (Licitacoes_Compras_e_Contratos/)"
    )
    parser.add_argument(
        "--output", type=Path, required=True,
        help="Diretorio destino (referencias_crie/)"
    )
    parser.add_argument(
        "--dry-run", action="store_true",
        help="Lista operacoes sem executar"
    )

    args = parser.parse_args()

    if not args.source.exists():
        print(f"ERRO: Diretorio fonte nao encontrado: {args.source}", file=sys.stderr)
        sys.exit(1)

    args.output.mkdir(parents=True, exist_ok=True)

    print(f"\n{'=' * 60}")
    print(f"CRIE Document Extraction")
    print(f"{'=' * 60}")
    print(f"Fonte:   {args.source}")
    print(f"Destino: {args.output}")
    print(f"Modo:    {'DRY RUN' if args.dry_run else 'EXECUCAO'}")
    print(f"Arquivos: {len(FILE_CONFIGS)}")
    print(f"{'=' * 60}\n")

    results = []
    for config in FILE_CONFIGS:
        result = process_file(config, args.source, args.output, args.dry_run)
        results.append(result)

    # Generate manifest and index
    if not args.dry_run:
        generate_manifest(results, args.output)
        generate_index_md(results, args.output)

    # Summary
    success = sum(1 for r in results if r["status"] == "success")
    failed = sum(1 for r in results if not r["status"].startswith(("success", "dry_run")))
    total_chars = sum(r.get("chars_extracted", 0) for r in results)

    print(f"\n{'=' * 60}")
    print(f"RESULTADO")
    print(f"{'=' * 60}")
    print(f"Total:       {len(results)}")
    print(f"Sucesso:     {success}")
    print(f"Falha:       {failed}")
    print(f"Caracteres:  {total_chars:,}")
    print(f"{'=' * 60}\n")

    if failed > 0:
        print("Arquivos com problema:")
        for r in results:
            if r["status"] not in ("success", "dry_run"):
                print(f"  [{r['status']}] {r['src']}")
        sys.exit(1)
    else:
        print("Todos os arquivos processados com sucesso!")


if __name__ == "__main__":
    main()
