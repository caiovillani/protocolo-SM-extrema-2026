# src/context_engine/exporter.py

"""Export functionality for Context Engine command outputs.

This module provides export capabilities for command outputs in multiple formats:
- Markdown (MD)
- YAML
- JSON
- DOCX (Word document)

All user-facing text is in Portuguese (Brazilian).
"""

import json
import re
from dataclasses import dataclass
from datetime import datetime
from enum import Enum
from io import BytesIO
from pathlib import Path
from typing import Any, Dict, List, Optional, Union

import yaml

# DOCX import with graceful fallback
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


# ---------------------------------------------------------------------------
# Enums and Data Classes
# ---------------------------------------------------------------------------

class ExportFormat(Enum):
    """Supported export formats."""
    MARKDOWN = "md"
    YAML = "yaml"
    JSON = "json"
    DOCX = "docx"

    @classmethod
    def from_string(cls, value: str) -> "ExportFormat":
        """Create ExportFormat from string value.

        Args:
            value: Format string (e.g., "md", "yaml", "json", "docx")

        Returns:
            Corresponding ExportFormat enum

        Raises:
            ValueError: If format string is not recognized
        """
        value_lower = value.lower().strip()
        # Handle aliases
        aliases = {
            "markdown": "md",
            "yml": "yaml",
            "word": "docx",
            "doc": "docx",
        }
        value_lower = aliases.get(value_lower, value_lower)

        for fmt in cls:
            if fmt.value == value_lower:
                return fmt
        raise ValueError(f"Formato não reconhecido: '{value}'")

    @classmethod
    def list_formats(cls) -> List[str]:
        """List all available format values.

        Returns:
            List of format strings
        """
        return [fmt.value for fmt in cls]


@dataclass
class ExportResult:
    """Result of an export operation."""
    format: ExportFormat
    success: bool
    file_path: Optional[Path] = None
    size_bytes: int = 0
    timestamp: str = ""
    error_message: Optional[str] = None

    def __post_init__(self):
        if not self.timestamp:
            self.timestamp = datetime.now().isoformat()


@dataclass
class ExportMetadata:
    """Metadata included in exports."""
    command_name: str
    timestamp: str
    source: str = "Context Engine"
    version: str = "1.0"

    def to_dict(self) -> Dict[str, Any]:
        """Convert metadata to dictionary.

        Returns:
            Metadata dictionary
        """
        return {
            "comando": self.command_name,
            "data_hora": self.timestamp,
            "origem": self.source,
            "versao": self.version,
        }


# ---------------------------------------------------------------------------
# Export Parser Utilities
# ---------------------------------------------------------------------------

def parse_box_header(line: str) -> Optional[str]:
    """Parse box-drawing header line to extract title.

    Args:
        line: Line that may contain a header title

    Returns:
        Title text if found, None otherwise
    """
    # Check if line starts with spaces (indented title)
    stripped = line.strip()
    if stripped and not stripped.startswith("═") and not stripped.startswith("─"):
        return stripped
    return None


def parse_formatted_content(content: str) -> Dict[str, Any]:
    """Parse formatted console output into structured data.

    Args:
        content: Formatted console output string

    Returns:
        Dictionary with parsed sections and data
    """
    result: Dict[str, Any] = {
        "titulo": "",
        "secoes": [],
        "conteudo_raw": content,
    }

    lines = content.split("\n")
    current_section: Optional[Dict[str, Any]] = None
    in_header = False

    for i, line in enumerate(lines):
        # Detect header
        if line.startswith("═" * 10):
            in_header = True
            continue

        # Get title from header
        if in_header:
            if line.startswith("═" * 10):
                in_header = False
            else:
                title = parse_box_header(line)
                if title:
                    result["titulo"] = title
            continue

        # Detect section starts (lines with emoji followed by text and colon)
        emoji_pattern = r"^([^\w\s])\s+(.+):$"
        section_match = re.match(emoji_pattern, line.strip())
        if section_match and not line.strip().startswith("•"):
            if current_section:
                result["secoes"].append(current_section)
            current_section = {
                "emoji": section_match.group(1),
                "titulo": section_match.group(2),
                "linhas": [],
            }
            continue

        # Add content to current section
        if current_section and line.strip():
            current_section["linhas"].append(line)

    # Add final section
    if current_section:
        result["secoes"].append(current_section)

    return result


# ---------------------------------------------------------------------------
# CommandExporter Class
# ---------------------------------------------------------------------------

class CommandExporter:
    """Export command outputs in multiple formats."""

    def __init__(self):
        """Initialize the exporter."""
        self._check_docx_availability()

    def _check_docx_availability(self) -> None:
        """Check if DOCX support is available."""
        if not DOCX_AVAILABLE:
            # Log warning but don't fail
            pass

    def export_to_string(
        self,
        content: str,
        format: ExportFormat,
        metadata: Optional[ExportMetadata] = None,
    ) -> str:
        """Export content to string without writing to file.

        Args:
            content: Command output content to export
            format: Target export format
            metadata: Optional export metadata

        Returns:
            Exported content as string

        Raises:
            ValueError: If format requires file output (DOCX)
        """
        if format == ExportFormat.DOCX:
            raise ValueError(
                "Formato DOCX requer exportação para arquivo. "
                "Use export_to_file() em vez de export_to_string()."
            )

        if format == ExportFormat.MARKDOWN:
            return self._export_markdown(content, metadata)
        elif format == ExportFormat.YAML:
            return self._export_yaml(content, metadata)
        elif format == ExportFormat.JSON:
            return self._export_json(content, metadata)

        raise ValueError(f"Formato não suportado: {format}")

    def export_to_file(
        self,
        content: str,
        format: ExportFormat,
        output_path: Path,
        metadata: Optional[ExportMetadata] = None,
    ) -> ExportResult:
        """Export content to file.

        Args:
            content: Command output content to export
            format: Target export format
            output_path: Path for the output file
            metadata: Optional export metadata

        Returns:
            ExportResult with success status and details
        """
        try:
            # Validate path
            self._validate_output_path(output_path)

            # Create parent directories
            output_path.parent.mkdir(parents=True, exist_ok=True)

            # Export based on format
            if format == ExportFormat.DOCX:
                data = self._export_docx(content, metadata)
                output_path.write_bytes(data)
                size = len(data)
            else:
                exported = self.export_to_string(content, format, metadata)
                output_path.write_text(exported, encoding="utf-8")
                size = len(exported.encode("utf-8"))

            return ExportResult(
                format=format,
                success=True,
                file_path=output_path,
                size_bytes=size,
            )

        except PermissionError:
            return ExportResult(
                format=format,
                success=False,
                error_message=f"Sem permissão para escrever em: {output_path}",
            )
        except ValueError as e:
            return ExportResult(
                format=format,
                success=False,
                error_message=str(e),
            )
        except Exception as e:
            return ExportResult(
                format=format,
                success=False,
                error_message=f"Erro inesperado: {e}",
            )

    def _validate_output_path(self, path: Path) -> None:
        """Validate output path for security.

        Args:
            path: Path to validate

        Raises:
            ValueError: If path is invalid or potentially dangerous
        """
        # Check for path traversal
        try:
            resolved = path.resolve()
        except (OSError, ValueError) as e:
            raise ValueError(f"Caminho inválido: {e}")

        # Check for suspicious patterns
        path_str = str(path)
        if ".." in path_str:
            raise ValueError("Caminho inválido: não use '..'")

    def _export_markdown(
        self,
        content: str,
        metadata: Optional[ExportMetadata] = None,
    ) -> str:
        """Export as Markdown with metadata header.

        Args:
            content: Content to export
            metadata: Optional metadata

        Returns:
            Markdown formatted string
        """
        output_lines = []

        # Add metadata as YAML frontmatter if provided
        if metadata:
            output_lines.append("---")
            output_lines.append(f"comando: {metadata.command_name}")
            output_lines.append(f"data_hora: {metadata.timestamp}")
            output_lines.append(f"origem: {metadata.source}")
            output_lines.append("---")
            output_lines.append("")

        # Clean up box-drawing for Markdown
        # Convert ═══ headers to Markdown headers
        lines = content.split("\n")
        i = 0
        while i < len(lines):
            line = lines[i]

            # Skip box-drawing header lines
            if line.startswith("═" * 10):
                # Next non-empty line is the title
                i += 1
                if i < len(lines):
                    title = lines[i].strip()
                    if title and not title.startswith("═"):
                        output_lines.append(f"# {title}")
                        output_lines.append("")
                # Skip closing header line
                i += 1
                if i < len(lines) and lines[i].startswith("═"):
                    i += 1
                continue

            # Convert separator lines
            if line.startswith("─" * 10):
                output_lines.append("---")
                i += 1
                continue

            # Keep other lines as-is (emoji and content)
            output_lines.append(line)
            i += 1

        # Add footer
        output_lines.append("")
        output_lines.append("---")
        output_lines.append(
            f"*Exportado pelo Context Engine em "
            f"{datetime.now().strftime('%Y-%m-%d %H:%M')}*"
        )

        return "\n".join(output_lines)

    def _export_yaml(
        self,
        content: str,
        metadata: Optional[ExportMetadata] = None,
    ) -> str:
        """Export as YAML structured document.

        Args:
            content: Content to export
            metadata: Optional metadata

        Returns:
            YAML formatted string
        """
        parsed = parse_formatted_content(content)

        data: Dict[str, Any] = {
            "titulo": parsed["titulo"],
            "secoes": [],
            "exportacao": {
                "formato": "yaml",
                "data_hora": datetime.now().isoformat(),
                "origem": "Context Engine",
            },
        }

        if metadata:
            data["metadata"] = metadata.to_dict()

        # Convert sections
        for section in parsed["secoes"]:
            section_data = {
                "titulo": section["titulo"],
                "conteudo": "\n".join(section["linhas"]),
            }
            data["secoes"].append(section_data)

        # If no sections parsed, include raw content
        if not data["secoes"]:
            data["conteudo_raw"] = content

        return yaml.dump(
            data,
            allow_unicode=True,
            default_flow_style=False,
            sort_keys=False,
        )

    def _export_json(
        self,
        content: str,
        metadata: Optional[ExportMetadata] = None,
    ) -> str:
        """Export as JSON structured document.

        Args:
            content: Content to export
            metadata: Optional metadata

        Returns:
            JSON formatted string
        """
        parsed = parse_formatted_content(content)

        data: Dict[str, Any] = {
            "titulo": parsed["titulo"],
            "secoes": [],
            "exportacao": {
                "formato": "json",
                "data_hora": datetime.now().isoformat(),
                "origem": "Context Engine",
            },
        }

        if metadata:
            data["metadata"] = metadata.to_dict()

        # Convert sections
        for section in parsed["secoes"]:
            section_data = {
                "titulo": section["titulo"],
                "conteudo": "\n".join(section["linhas"]),
            }
            data["secoes"].append(section_data)

        # If no sections parsed, include raw content
        if not data["secoes"]:
            data["conteudo_raw"] = content

        return json.dumps(data, ensure_ascii=False, indent=2)

    def _export_docx(
        self,
        content: str,
        metadata: Optional[ExportMetadata] = None,
    ) -> bytes:
        """Export as DOCX (Word document).

        Args:
            content: Content to export
            metadata: Optional metadata

        Returns:
            DOCX file as bytes

        Raises:
            ImportError: If python-docx is not installed
        """
        if not DOCX_AVAILABLE:
            raise ImportError(
                "Biblioteca python-docx não está instalada.\n"
                "Instale com: pip install python-docx"
            )

        doc = Document()

        # Parse content
        parsed = parse_formatted_content(content)

        # Add title
        if parsed["titulo"]:
            heading = doc.add_heading(parsed["titulo"], level=1)
            # Style the heading
            for run in heading.runs:
                run.font.size = Pt(16)
                run.font.color.rgb = RGBColor(0, 51, 102)

        # Add metadata table if provided
        if metadata:
            doc.add_heading("Metadados", level=2)
            table = doc.add_table(rows=4, cols=2)
            table.style = "Table Grid"

            rows_data = [
                ("Comando", metadata.command_name),
                ("Data/Hora", metadata.timestamp),
                ("Origem", metadata.source),
                ("Versão", metadata.version),
            ]

            for i, (label, value) in enumerate(rows_data):
                table.cell(i, 0).text = label
                table.cell(i, 1).text = value
                # Bold the labels
                table.cell(i, 0).paragraphs[0].runs[0].bold = True

            doc.add_paragraph()  # Spacing

        # Add sections
        if parsed["secoes"]:
            for section in parsed["secoes"]:
                # Section heading
                doc.add_heading(section["titulo"], level=2)

                # Section content
                for line in section["linhas"]:
                    line_stripped = line.strip()

                    # Handle bullet points
                    if line_stripped.startswith("•"):
                        text = line_stripped[1:].strip()
                        doc.add_paragraph(text, style="List Bullet")
                    elif line_stripped.startswith(tuple("0123456789")):
                        # Numbered list
                        doc.add_paragraph(line_stripped, style="List Number")
                    elif line_stripped:
                        doc.add_paragraph(line_stripped)
        else:
            # No sections parsed, add raw content
            doc.add_heading("Conteúdo", level=2)
            for line in content.split("\n"):
                # Skip box-drawing lines
                if line.startswith("═") or line.startswith("─"):
                    continue
                line_stripped = line.strip()
                if line_stripped:
                    if line_stripped.startswith("•"):
                        text = line_stripped[1:].strip()
                        doc.add_paragraph(text, style="List Bullet")
                    else:
                        doc.add_paragraph(line_stripped)

        # Add footer
        doc.add_paragraph()
        doc.add_paragraph("─" * 40)
        footer = doc.add_paragraph(
            f"Exportado pelo Context Engine em "
            f"{datetime.now().strftime('%Y-%m-%d %H:%M')}"
        )
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in footer.runs:
            run.font.size = Pt(9)
            run.font.italic = True

        # Save to bytes
        buffer = BytesIO()
        doc.save(buffer)
        return buffer.getvalue()


# ---------------------------------------------------------------------------
# Convenience Functions
# ---------------------------------------------------------------------------

def export_command_output(
    command_name: str,
    output: str,
    format: Union[ExportFormat, str],
    output_path: Optional[Path] = None,
) -> ExportResult:
    """Convenience function to export command output.

    Args:
        command_name: Name of the command that generated output
        output: Command output to export
        format: Export format (ExportFormat or string)
        output_path: Optional output path (auto-generated if None)

    Returns:
        ExportResult with success status and details
    """
    # Convert string format to enum
    if isinstance(format, str):
        format = ExportFormat.from_string(format)

    # Create metadata
    metadata = ExportMetadata(
        command_name=command_name,
        timestamp=datetime.now().isoformat(),
    )

    # Auto-generate output path if not provided
    if output_path is None:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{command_name}_{timestamp}.{format.value}"
        output_path = Path("./exports") / filename

    # Export
    exporter = CommandExporter()
    return exporter.export_to_file(output, format, output_path, metadata)
